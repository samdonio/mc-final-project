import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def jsonToDoc(path):
    
    # Read JSON
    with open (path, 'r') as file:
        content = json.load(file)

    # Construct document
    document = Document()

    # Set defaults
    font = 'Times New Roman'

    # Write title and subtitle
    title = document.add_heading('Mobile Computing (CMSC 23400): Final Project'
                                 + '\n' + 'By Good Partner Solutions (GPS)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = font
    title.runs[0].font.bold = True

    subtitle = document.add_heading('Output', level=1)
    subtitle.runs[0].font.name = font

    # Write body
    body = document.add_paragraph(content)
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.runs[0].font.name = font

    document.save("txt.docx")